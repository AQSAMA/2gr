#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_production import (
    ASSEMBLED_DIR,
    FIGURES_DIR,
    METHOD_A_DIR,
    REPO_ROOT,
    assemble_markdown,
    copy_figure_assets,
    ensure_dirs as ensure_production_dirs,
    iter_markdown_blocks,
)

TYPST_CONTENT_DIR = REPO_ROOT / "typst_content"
TYPST_OUTPUT_DIR = TYPST_CONTENT_DIR / "output"
TYPST_SOURCE = TYPST_CONTENT_DIR / "research.typ"
TYPST_PDF = TYPST_OUTPUT_DIR / "research.pdf"
TYPST_DOCX = TYPST_OUTPUT_DIR / "research.docx"

TITLE = "Psychiatric Medication Use and Public Acceptance in Iraq"
# Student names are sorted alphabetically for every downstream use
# (cover, supervisor certification, DOCX metadata).
STUDENTS = sorted([
    "Abdul Rahman Wakaa Ali",
    "Ali Basem Hammoud",
    "Shifa Safi Aboud",
    "Zainab Mashal Nayef",
])
SUPERVISOR = "Hameed Adnan"
UNIVERSITY = "University of Al-Maarif"
COLLEGE = "College of Pharmacy"
MONTH_YEAR = "May, 2026"
LOGO_CANDIDATES = (
    "University_logo.png",
    "university logo.png",
    "university_logo.png",
    "University logo.png",
    "University Logo.png",
    "almaarif logo.png",
    "al-maarif logo.png",
)

# Muted academic navy for in-text citation highlighting. Subtle enough to
# preserve readability while distinguishing citations from body prose.
CITATION_COLOR_HEX = "2C5282"  # used by python-docx RGBColor.from_string
CITATION_COLOR_TYPST = "#2c5282"

# A parenthetical in-text citation starts with an uppercase author surname
# and contains a 4-digit year. This avoids matching numeric parentheticals
# like "(0.504)", statistical inline stats like "(p=0.0001)", or reference
# year suffixes like "(2020): 56" that live inside the bibliography lines.
CITATION_REGEX = re.compile(r"\([A-Z][^()]*?\d{4}[a-z]?\)")

# Pairs of (surname prefix, distinctive title keyword) that identify each
# reference actually cited in the manuscript text. Used to filter the master
# bibliography so only cited sources appear in the final rendered list.
CITED_REFERENCES: list[tuple[str, str]] = [
    ("Abdisa", "Self-Stigma"),
    ("Angermeyer", "Public Attitudes towards Psychiatry"),
    ("Booth", "Conundrum for Healthcare"),
    ("Burnam", "War Veterans"),
    ("Cubillos", "Integrating Mental Health Services"),
    ("Elyamani", "Arab States of the Gulf"),
    ("Gast", "Medication Adherence Influencing"),
    ("Horne", "Beliefs about Medicines"),
    ("Hussein Alwan", "Stigma Toward Mental Illness"),
    ("Kadhim", "ANXIETY, DEPRESSION"),
    ("Kalaman", "Parental Factors"),
    ("Mojtabai", "Americans"),
    ("Nassr", "High-Dose Antipsychotic"),
    ("Okasha", "Middle East, and North Africa"),
    ("Rasheed", "Health Care-Seeking"),
    ("Sadik", "Public Perception of Mental Health"),
    ("Saied", "Mental Health Crisis"),
    ("Slewa-Younan", "Resettled Iraqi Refugees"),
    ("Younis", "Personal History of Psychiatry"),
    ("Younis", "Faith Healers"),
    ("Zolezzi", "Stigma Associated with Mental Illness"),
    ("Zygmunt", "Interventions to Improve Medication Adherence"),
]


def find_university_logo() -> Path | None:
    search_dirs = (REPO_ROOT, FIGURES_DIR)
    for directory in search_dirs:
        for filename in LOGO_CANDIDATES:
            candidate = directory / filename
            if candidate.exists():
                return candidate
    for directory in search_dirs:
        for candidate in sorted(directory.glob("*.png")):
            lowered = candidate.name.lower()
            if "logo" in lowered or "university" in lowered or "maarif" in lowered:
                return candidate
    return None


def typst_string(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def clean_text(value: str) -> str:
    value = value.replace("\u00a0", " ")
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", value)
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def ensure_dirs() -> None:
    TYPST_CONTENT_DIR.mkdir(parents=True, exist_ok=True)
    TYPST_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def is_cited_reference(line: str) -> bool:
    """Return True if a bibliography line matches at least one cited source."""
    for surname, keyword in CITED_REFERENCES:
        # Surname should appear near the very start of the reference entry.
        head = line[: len(surname) + 6]
        if surname in head and keyword.lower() in line.lower():
            return True
    return False


def filter_cited_references(references: list[str]) -> list[str]:
    filtered = [ref for ref in references if is_cited_reference(ref)]
    # Preserve the master bibliography ordering (alphabetical) but remove
    # duplicates that the filter may introduce when keyword matches overlap.
    seen = set()
    unique: list[str] = []
    for ref in filtered:
        if ref not in seen:
            seen.add(ref)
            unique.append(ref)
    return unique


def _split_references(md_text: str) -> tuple[str, list[str]]:
    marker = "\n# VIII. REFERENCES\n"
    if marker not in md_text:
        return md_text, []
    before, after = md_text.split(marker, 1)
    references = [clean_text(line) for line in after.splitlines() if clean_text(line)]
    return before + marker, filter_cited_references(references)


def collect_manuscript_calls(md_path: Path) -> tuple[list[str], list[str]]:
    """Return Typst calls for roman-numbered front matter and arabic main matter."""
    front_calls: list[str] = []
    main_calls: list[str] = []
    target = front_calls
    skip_cover = True
    in_references = False
    md_text, references = _split_references(md_path.read_text(encoding="utf-8"))

    for kind, data in iter_markdown_blocks(md_text):
        if skip_cover:
            if kind == "h1" and data.strip().upper() == "ABSTRACT":
                skip_cover = False
                target = main_calls
                target.append("#start-main-numbering()")
                target.append('#set-running-head("")')
            else:
                continue

        if kind == "chaptertitle":
            chapter_number, chapter_name = (data.split("|||", 1) + [""])[:2]
            target.append(f"#chapter-page({typst_string(chapter_number)}, {typst_string(chapter_name)})")
            target.append(f"#set-running-head({typst_string(chapter_name)})")
            continue

        if kind == "h1":
            text = clean_text(data)
            if text.upper() == "VIII. REFERENCES":
                in_references = True
                # References live outside any numbered chapter, so switch the
                # running head to "References" before emitting the heading.
                target.append('#set-running-head("References")')
            target.append(f"#section-title({typst_string(text)})")
            continue

        if kind == "h2":
            target.append(f"#h2({typst_string(clean_text(data))})")
            continue

        if kind == "paragraph":
            text = clean_text(data)
            if not text:
                continue
            fn = "refp" if in_references else "p"
            target.append(f"#{fn}({typst_string(text)})")
            continue

        if kind == "image":
            caption, rel_path = data.split("|||", 1)
            filename = Path(rel_path).name
            target.append(f"#fig({typst_string('../figures/' + filename)}, {typst_string(clean_text(caption))})")
            continue

        if kind == "pagebreak":
            target.append("#pagebreak()")
            continue

    if references:
        if not in_references:
            target.append('#set-running-head("References")')
            target.append(f"#section-title({typst_string('VIII. REFERENCES')})")
        for reference in references:
            target.append(f"#refp({typst_string(reference)})")

    return front_calls, main_calls


def render_typst_source(md_path: Path) -> str:
    front_calls, main_calls = collect_manuscript_calls(md_path)
    students = ", ".join(STUDENTS)
    student_lines = "\\\n".join(STUDENTS)

    logo_path = find_university_logo()
    logo_block = ""
    if logo_path is not None:
        logo_rel = Path("..") / logo_path.relative_to(REPO_ROOT)
        logo_block = f"  #image({typst_string(logo_rel.as_posix())}, width: 2.25cm)\n  #v(0.16cm)\n"

    preamble = f'''// Editable Typst source for the graduation project.
// Generated from content/*.md by production/src/build_typst_content.py.
// The design follows common university thesis conventions: bordered A4 pages,
// formal title page, compact front matter, automatic contents, and figure lists.

#set document(title: {typst_string(TITLE)}, author: {typst_string(students)})

#let navy = rgb("#102a43")
#let gold = rgb("#b58b2a")
#let ink = rgb("#111827")
#let pale = rgb("#f7f9fc")
#let citation-color = rgb("{CITATION_COLOR_TYPST}")
#let page-border = rect(width: 100%, height: 100%, stroke: 0.8pt + navy)
#let running-head = state("running-head", "")
#let set-running-head(s) = running-head.update(s)

// The running head sits at the top RIGHT and is suppressed on the first
// page of every chapter. Chapter-title pages already disable the header
// via `set page(header: none)`; content pages mark the opening page with a
// `<section-start>` label so this context block knows to hide the head on
// that page only.
#let regular-page-header = context {{
  let head = running-head.get()
  if head != "" {{
    let page-num = here().page()
    let markers = query(<section-start>).filter(m => m.location().page() == page-num)
    if markers.len() == 0 {{
      align(right)[#text(size: 9pt, fill: navy, style: "italic")[#head]]
    }}
  }}
}}

#set page(paper: "a4", margin: 1.5cm, background: page-border, header: regular-page-header)
#set text(font: ("Times New Roman", "Times"), size: 14pt, fill: ink)
#set par(leading: 0.55em, justify: true)

// Muted academic navy for parenthetical in-text citations such as
// "(Author, 2020)" or "(Author et al., 2020; Other, 2019)". The leading
// capital letter in the required lookbehind avoids matching numeric
// parentheses like "(p=0.0001)" or "(adjusted OR 0.504, ...)".
#show regex("\\([A-Z][^()]*?\\d{{4}}[a-z]?\\)"): it => text(fill: citation-color, it)

#show heading.where(level: 1): it => block(above: 10pt, below: 8pt)[
  #text(size: 18pt, weight: "bold", fill: navy)[#it.body]
  #linebreak()
  #line(length: 100%, stroke: 0.7pt + gold)
]
#show heading.where(level: 2): it => block(above: 8pt, below: 6pt)[
  #text(size: 16pt, weight: "bold", fill: navy)[#it.body]
]
#show figure: it => block(above: 10pt, below: 12pt, inset: 6pt, stroke: 0.35pt + rgb("#c9d4e5"))[#align(center)[#it]]

#let center-line(s, size: 14pt, weight: "regular", fill: ink) = align(center)[#text(size: size, weight: weight, fill: fill)[#s]]
#let p(s) = par(first-line-indent: 1.27cm, justify: true)[#s]
#let refp(s) = block(above: 3pt, below: 5pt)[
  #set text(size: 12pt)
  #par(first-line-indent: 0pt, hanging-indent: 0.5in, justify: true)[#s]
]
#let h1(s) = heading(level: 1, outlined: true)[#s]
#let section-title(s) = [
  #pagebreak(weak: true)
  #[#metadata(none) <section-start>]
  #h1(s)
]
#let h2(s) = heading(level: 2, outlined: true)[#s]
#let fig(path, caption-text) = figure(image(path, width: 90%), caption: [#caption-text])

#let front-title(s) = [
  #align(center)[
    #box(width: 80%, inset: 10pt, stroke: 0.7pt + gold, fill: pale)[
      #text(size: 22pt, weight: "bold", fill: navy)[#s]
    ]
  ]
]

#let chapter-page(chapter, title) = [
  #pagebreak(weak: true)
  #set page(header: none)
  #align(center + horizon)[
    #box(width: 84%, inset: 28pt, stroke: 1pt + navy, fill: pale)[
      #align(center)[
        #text(size: 32pt, weight: "bold", fill: navy)[#chapter]
        #v(8pt)
        #line(length: 55%, stroke: 0.8pt + gold)
        #v(8pt)
        #text(size: 20pt, weight: "bold", fill: navy)[#title]
      ]
    ]
  ]
  #pagebreak()
  #set page(header: regular-page-header)
]

#let start-main-numbering() = [
  #pagebreak(weak: true)
  #set page(numbering: "1", number-align: top + center)
  #counter(page).update(1)
]

// Cover page: unnumbered. Certification begins on the second page.
#set page(numbering: none)
#align(center)[
{logo_block}  #text(size: 15pt, weight: "bold", fill: navy)[Republic of Iraq] \\
  #text(size: 15pt, weight: "bold", fill: navy)[Ministry of Higher Education and Scientific Research] \\
  #text(size: 15pt, weight: "bold", fill: navy)[{UNIVERSITY}] \\
  #text(size: 15pt, weight: "bold", fill: navy)[{COLLEGE}]
  #v(0.35cm)
  #box(width: 88%, inset: 12pt, stroke: 1pt + gold, fill: pale)[
    #text(size: 24pt, weight: "bold", fill: navy)[{TITLE}]
  ]
  #v(0.32cm)
  #text(size: 14pt)[A Project Submitted to] \\
  #text(size: 13pt)[The {COLLEGE}, {UNIVERSITY}, Department of Clinical Pharmacy, in Partial Fulfillment for the Bachelor of Pharmacy]
  #v(0.28cm)
  #text(size: 14pt, weight: "bold")[By] \\
  #text(size: 20pt, weight: "bold", fill: navy)[{student_lines}]
  #v(0.25cm)
  #text(size: 14pt, weight: "bold")[Supervised by:] \\
  #text(size: 20pt, weight: "bold", fill: navy)[{SUPERVISOR}] \\
  #text(size: 16pt)[Supervisor's Degree]
  #v(0.18cm)
  #text(size: 14pt)[{MONTH_YEAR}]
]

// Roman-numbered preliminary pages. Numbering stays Roman through every
// front-matter spread and switches to Arabic the moment the abstract opens.
#pagebreak()
#set page(numbering: "i", number-align: top + center)
#counter(page).update(1)
#front-title[Certification of the Supervisor]
#p("I certify that this project entitled \u201c{TITLE}\u201d was prepared by the fifth-year students {students} under my supervision at the {COLLEGE}/{UNIVERSITY} in partial fulfillment of the graduation requirements for the Bachelor Degree in Pharmacy.")
#align(right)[#text(weight: "bold")[Supervisor's name: {SUPERVISOR}]]
#v(0.55cm)
#front-title[Dedication]
#p("We dedicate this work to our families, whose patience made long study days easier, and to every Iraqi patient who deserves safe, respectful, and evidence-based mental health care. We also dedicate it to the teachers and pharmacists who taught us that science becomes meaningful when it serves people with honesty and compassion.")
#v(0.35cm)
#front-title[Acknowledgment]
#p("We thank Dr. {SUPERVISOR} for his supervision, guidance, and careful advice throughout this project. We are also grateful to the College of Pharmacy at {UNIVERSITY}, to the participants who gave their time to answer the survey, and to our colleagues who supported the data collection and revision process.")

#pagebreak()
#front-title[Table of Contents]
// The table of contents is tightened to fit a single page. We keep depth 2
// so readers still see chapter + section structure, but compact the entry
// size and spacing so the full outline renders inside one A4 page.
#{{
  show outline.entry: it => block(above: 0.15em, below: 0.15em, text(size: 10pt, it))
  outline(title: none, depth: 2, indent: auto)
}}

#pagebreak()
#front-title[List of Figures]
#outline(title: none, target: figure.where(kind: image))
#v(0.4cm)
#front-title[List of Tables]
#p("No manuscript tables are currently embedded as formal Typst tables in this editable source. Statistical results are reported in the text and figures.")
#v(0.4cm)
#front-title[List of Abbreviations]
#par(first-line-indent: 0pt)[AOR: Adjusted Odds Ratio \\
CI: Confidence Interval \\
LLR: Likelihood Ratio Test \\
MLE: Maximum Likelihood Estimation \\
OR: Odds Ratio \\
PTSD: Post-Traumatic Stress Disorder \\
RRR: Relative Risk Ratio \\
Q6/Q7/Q8/Q9/Q11/Q12/Q13: Survey question item codes used in analysis and reporting \\
R\u00b2: Coefficient of determination, reported as pseudo R\u00b2 in logistic model fit summaries]

'''
    return preamble + "\n".join(front_calls) + "\n\n" + "\n".join(main_calls) + "\n"


def write_typst_source(md_path: Path) -> Path:
    source = render_typst_source(md_path)
    TYPST_SOURCE.write_text(source, encoding="utf-8")
    return TYPST_SOURCE


def _print_process_output(output: str | bytes | None) -> None:
    if output is None:
        return
    if isinstance(output, bytes):
        output = output.decode(errors="replace")
    if output.strip():
        print(output.strip())


def compile_typst_pdf() -> bool:
    typst = shutil.which("typst")
    if typst is None:
        print("WARNING: typst is not installed; typst_content PDF compilation was skipped.")
        return False
    result = subprocess.run(
        [typst, "compile", "--root", str(REPO_ROOT), str(TYPST_SOURCE), str(TYPST_PDF)],
        cwd=REPO_ROOT,
        text=True,
        capture_output=True,
        check=False,
        timeout=300,
    )
    if result.returncode != 0:
        print("WARNING: Typst failed for typst_content/research.typ; continuing with other outputs.")
        _print_process_output(result.stdout)
        _print_process_output(result.stderr)
        return False
    print(f"Typst content PDF: {TYPST_PDF}")
    return True


def collect_docx_blocks(md_path: Path) -> list[tuple[str, str]]:
    """Collect the same manuscript structure used by the editable Typst source."""
    blocks: list[tuple[str, str]] = []
    skip_cover = True
    in_references = False
    md_text, references = _split_references(md_path.read_text(encoding="utf-8"))

    for kind, data in iter_markdown_blocks(md_text):
        if skip_cover:
            if kind == "h1" and data.strip().upper() == "ABSTRACT":
                skip_cover = False
                blocks.append(("start_main", ""))
            else:
                continue

        if kind == "chaptertitle":
            chapter_number, chapter_name = (data.split("|||", 1) + [""])[:2]
            blocks.append(("chapter", f"{chapter_number}|||{chapter_name}"))
            continue

        if kind == "h1":
            text = clean_text(data)
            if text.upper() == "VIII. REFERENCES":
                in_references = True
                blocks.append(("references_section", text))
                continue
            blocks.append(("section", text))
            continue

        if kind == "h2":
            blocks.append(("h2", clean_text(data)))
            continue

        if kind == "paragraph":
            text = clean_text(data)
            if text:
                blocks.append(("reference" if in_references else "paragraph", text))
            continue

        if kind == "image":
            caption, rel_path = data.split("|||", 1)
            blocks.append(("image", f"{clean_text(caption)}|||{Path(rel_path).name}"))
            continue

        if kind == "pagebreak":
            blocks.append(("pagebreak", ""))
            continue

    if references:
        if not in_references:
            blocks.append(("references_section", "VIII. REFERENCES"))
        for reference in references:
            blocks.append(("reference", reference))

    return blocks


def _set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_border(paragraph, color: str = "B58B2A", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), color)
        borders.append(element)


def _set_page_border(section) -> None:
    sect_pr = section._sectPr
    borders = sect_pr.find(qn("w:pgBorders"))
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "18")
        element.set(qn("w:color"), "102A43")
        borders.append(element)


def _set_page_numbering(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def _add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    result = OxmlElement("w:t")
    result.text = " "
    run._r.append(result)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _configure_section(
    section,
    *,
    numbered: bool,
    number_format: str = "decimal",
    start: int | None = 1,
    running_head: str = "",
    empty_first_running_head: bool = False,
) -> None:
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    _set_page_border(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    # Academic convention: the first page of each chapter shows no running
    # head, so we enable Word's "different first page" whenever the section
    # has a running head to suppress on its opener.
    section.different_first_page_header_footer = empty_first_running_head

    for paragraph in section.header.paragraphs:
        paragraph.clear()
    for paragraph in section.first_page_header.paragraphs:
        paragraph.clear()
    for paragraph in section.footer.paragraphs:
        paragraph.clear()
    for paragraph in section.first_page_footer.paragraphs:
        paragraph.clear()

    if numbered:
        if start is not None:
            _set_page_numbering(section, number_format, start)

        # Centered page number in the footer keeps pagination clean and
        # independent from the running head in the header.
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.first_line_indent = Inches(0)
        _add_field_run(footer, "PAGE")
        if empty_first_running_head:
            first_footer = section.first_page_footer.paragraphs[0]
            first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_footer.paragraph_format.first_line_indent = Inches(0)
            _add_field_run(first_footer, "PAGE")

        if running_head:
            # Running head lives in the header, right-aligned, italic. The
            # first-page header of the section stays empty so chapter-opening
            # pages are clean per academic convention.
            header = section.header.paragraphs[0]
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header.paragraph_format.first_line_indent = Inches(0)
            run = header.add_run(running_head)
            run.italic = True
            _set_run_font(run, size=9, color="102A43")


def _setup_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 16)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("102A43")

    # Compact TOC styles help the automatic Word TOC fit on one page when
    # users refresh fields after opening the document.
    for style_name, size in [("TOC 1", 11), ("TOC 2", 10), ("TOC 3", 10)]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(size)


def _center_paragraph(doc: Document, text: str = "", size: int | float = 14, bold: bool = False, color: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Inches(0)
    if text:
        run = paragraph.add_run(text)
        _set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def _front_title(doc: Document, title: str) -> None:
    paragraph = _center_paragraph(doc, title, size=22, bold=True, color="102A43")
    _set_paragraph_border(paragraph)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)


def _add_paragraph_with_citations(doc: Document, text: str, *, size: float = 14, hanging: bool = False, justify: bool = True) -> None:
    """Add a body paragraph, colouring parenthetical in-text citations."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    if hanging:
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(5)
    else:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cursor = 0
    for match in CITATION_REGEX.finditer(text):
        if match.start() > cursor:
            plain_run = paragraph.add_run(text[cursor:match.start()])
            _set_run_font(plain_run, size=size)
        cite_run = paragraph.add_run(match.group(0))
        _set_run_font(cite_run, size=size, color=CITATION_COLOR_HEX)
        cursor = match.end()
    if cursor < len(text):
        tail = paragraph.add_run(text[cursor:])
        _set_run_font(tail, size=size)


def _add_body_paragraph(doc: Document, text: str) -> None:
    _add_paragraph_with_citations(doc, text, size=14, hanging=False)


def _add_reference_paragraph(doc: Document, text: str) -> None:
    # Reference entries use a hanging indent and a smaller body size.
    # Citation coloring is intentionally skipped here because the Chicago
    # bibliography format uses "(Year)" which is not an in-text citation.
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        _set_run_font(run, size=12)


def _add_cover_page(doc: Document) -> None:
    logo_path = find_university_logo()
    if logo_path is not None:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Inches(0)
        run = paragraph.add_run()
        run.add_picture(str(logo_path), width=Inches(0.9))
    for line in ["Republic of Iraq", "Ministry of Higher Education and Scientific Research", UNIVERSITY, COLLEGE]:
        _center_paragraph(doc, line, size=15, bold=True, color="102A43")
    title = _center_paragraph(doc, TITLE, size=24, bold=True, color="102A43")
    _set_paragraph_border(title)
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(14)
    _center_paragraph(doc, "A Project Submitted to", size=14)
    _center_paragraph(
        doc,
        f"The {COLLEGE}, {UNIVERSITY}, Department of Clinical Pharmacy, in Partial Fulfillment for the Bachelor of Pharmacy",
        size=13,
    )
    doc.add_paragraph()
    _center_paragraph(doc, "By", size=14, bold=True)
    for student in STUDENTS:
        _center_paragraph(doc, student, size=20, bold=True, color="102A43")
    doc.add_paragraph()
    _center_paragraph(doc, "Supervised by:", size=14, bold=True)
    _center_paragraph(doc, SUPERVISOR, size=20, bold=True, color="102A43")
    _center_paragraph(doc, "Supervisor's Degree", size=16)
    doc.add_paragraph()
    _center_paragraph(doc, MONTH_YEAR, size=14)


def _add_preliminary_pages(doc: Document, figure_captions: list[str]) -> None:
    _front_title(doc, "Certification of the Supervisor")
    _add_body_paragraph(
        doc,
        f"I certify that this project entitled \u201c{TITLE}\u201d was prepared by the fifth-year students {', '.join(STUDENTS)} under my supervision at the {COLLEGE}/{UNIVERSITY} in partial fulfillment of the graduation requirements for the Bachelor Degree in Pharmacy.",
    )
    paragraph = doc.add_paragraph(f"Supervisor's name: {SUPERVISOR}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.runs[0].bold = True
    _set_run_font(paragraph.runs[0], size=14, bold=True)

    _front_title(doc, "Dedication")
    _add_body_paragraph(
        doc,
        "We dedicate this work to our families, whose patience made long study days easier, and to every Iraqi patient who deserves safe, respectful, and evidence-based mental health care. We also dedicate it to the teachers and pharmacists who taught us that science becomes meaningful when it serves people with honesty and compassion.",
    )
    _front_title(doc, "Acknowledgment")
    _add_body_paragraph(
        doc,
        f"We thank Dr. {SUPERVISOR} for his supervision, guidance, and careful advice throughout this project. We are also grateful to the {COLLEGE} at {UNIVERSITY}, to the participants who gave their time to answer the survey, and to our colleagues who supported the data collection and revision process.",
    )

    doc.add_page_break()
    _front_title(doc, "Table of Contents")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    # Depth 2 is enough to show chapters + main sections and, combined with
    # the tightened TOC styles in _setup_docx_styles, keeps the outline
    # compact enough to fit on a single page when Word refreshes fields.
    _add_field_run(paragraph, r'TOC \o "1-2" \h \z \u')

    doc.add_page_break()
    _front_title(doc, "List of Figures")
    for caption in figure_captions:
        paragraph = doc.add_paragraph(caption)
        paragraph.paragraph_format.first_line_indent = Inches(0)
    _front_title(doc, "List of Tables")
    _add_body_paragraph(doc, "No manuscript tables are currently embedded as formal tables in this editable source. Statistical results are reported in the text and figures.")
    _front_title(doc, "List of Abbreviations")
    abbreviations = [
        "AOR: Adjusted Odds Ratio",
        "CI: Confidence Interval",
        "LLR: Likelihood Ratio Test",
        "MLE: Maximum Likelihood Estimation",
        "OR: Odds Ratio",
        "PTSD: Post-Traumatic Stress Disorder",
        "RRR: Relative Risk Ratio",
        "Q6/Q7/Q8/Q9/Q11/Q12/Q13: Survey question item codes used in analysis and reporting",
        "R\u00b2: Coefficient of determination, reported as pseudo R\u00b2 in logistic model fit summaries",
    ]
    for item in abbreviations:
        paragraph = doc.add_paragraph(item)
        paragraph.paragraph_format.first_line_indent = Inches(0)


def build_typst_content_docx(md_path: Path, out_path: Path) -> None:
    blocks = collect_docx_blocks(md_path)
    figure_captions = [data.split("|||", 1)[0] for kind, data in blocks if kind == "image"]

    doc = Document()
    # Store the sorted student list in the document metadata so tools that
    # read DOCX properties see the same alphabetical ordering used visually.
    doc.core_properties.title = TITLE
    doc.core_properties.author = ", ".join(STUDENTS)

    _setup_docx_styles(doc)
    _configure_section(doc.sections[0], numbered=False)
    _add_cover_page(doc)

    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(front_section, numbered=True, number_format="lowerRoman", start=1)
    _add_preliminary_pages(doc, figure_captions)

    main_started = False
    in_references = False
    chapter_just_added = False
    current_running_head = ""

    def emit_section_page_break() -> None:
        nonlocal chapter_just_added
        if chapter_just_added:
            chapter_just_added = False
        else:
            doc.add_page_break()

    for kind, data in blocks:
        if kind == "start_main" and not main_started:
            # The abstract is the first Arabic-numbered page. It opens a new
            # Word section with decimal numbering starting at 1 and no
            # running head (academic convention keeps front-matter-like
            # items such as the abstract un-titled in the header).
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            _configure_section(section, numbered=True, number_format="decimal", start=1)
            main_started = True
            chapter_just_added = False
            current_running_head = ""
            continue

        if kind == "chapter":
            chapter_number, chapter_name = (data.split("|||", 1) + [""])[:2]
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            _configure_section(
                section,
                numbered=True,
                number_format="decimal",
                start=None,
                running_head=chapter_name,
                empty_first_running_head=True,
            )
            current_running_head = chapter_name
            paragraph = _center_paragraph(doc, chapter_number, size=32, bold=True, color="102A43")
            paragraph.paragraph_format.space_before = Inches(3)
            _center_paragraph(doc, chapter_name, size=20, bold=True, color="102A43")
            doc.add_page_break()
            chapter_just_added = True
            continue

        if kind == "references_section":
            # References get a dedicated section so the running head reads
            # "References" rather than inheriting the previous chapter.
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            _configure_section(
                section,
                numbered=True,
                number_format="decimal",
                start=None,
                running_head="References",
                empty_first_running_head=True,
            )
            current_running_head = "References"
            paragraph = doc.add_paragraph(data)
            paragraph.style = doc.styles["Heading 1"]
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_after = Pt(8)
            in_references = True
            chapter_just_added = True
            continue

        if kind == "section":
            emit_section_page_break()
            paragraph = doc.add_paragraph(data)
            paragraph.style = doc.styles["Heading 1"]
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_after = Pt(8)
            in_references = data.upper() == "VIII. REFERENCES"
            continue

        if kind == "h2":
            paragraph = doc.add_paragraph(data)
            paragraph.style = doc.styles["Heading 2"]
            paragraph.paragraph_format.first_line_indent = Inches(0)
            continue

        if kind == "paragraph":
            _add_body_paragraph(doc, data)
            continue

        if kind == "reference":
            _add_reference_paragraph(doc, data)
            continue

        if kind == "image":
            caption, filename = data.split("|||", 1)
            image_path = FIGURES_DIR / filename
            if image_path.exists():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Inches(0)
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(5.9))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = Inches(0)
            for run in cap.runs:
                _set_run_font(run, size=12, bold=True)
            continue

        if kind == "pagebreak":
            doc.add_page_break()
            chapter_just_added = False

    doc.save(out_path)
    print(f"Typst content DOCX: {out_path}")


def copy_docx_output() -> bool:
    source_docx = METHOD_A_DIR / "research_method_a.docx"
    if not source_docx.exists():
        print("WARNING: Method A DOCX was not found; typst_content DOCX fallback was skipped.")
        return False
    shutil.copy2(source_docx, TYPST_DOCX)
    print(f"Typst content DOCX emergency fallback copied from Method A: {TYPST_DOCX}")
    return True


def run_typst_content(md_path: Path | None = None) -> None:
    ensure_dirs()
    if md_path is None:
        ensure_production_dirs()
        copy_figure_assets()
        md_path = ASSEMBLED_DIR / "comprehensive_research.md"
        if not md_path.exists():
            md_path = assemble_markdown()
    if find_university_logo() is None:
        print("WARNING: university logo PNG was not found at repository root or in figures/; cover logo placement was skipped.")
    source_path = write_typst_source(md_path)
    print(f"Editable Typst source: {source_path}")
    compile_typst_pdf()
    try:
        build_typst_content_docx(md_path, TYPST_DOCX)
    except Exception as exc:
        print(f"WARNING: Typst content DOCX generation failed ({exc}); falling back if possible.")
        copy_docx_output()


def main() -> None:
    run_typst_content()


if __name__ == "__main__":
    main()
