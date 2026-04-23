#!/usr/bin/env python3
from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
import markdown2
from xhtml2pdf import pisa
import pypandoc


REPO_ROOT = Path(__file__).resolve().parents[2]
PROD_ROOT = REPO_ROOT / "production"
CONTENT_DIR = REPO_ROOT / "content"
FIGURES_DIR = REPO_ROOT / "figures"
REFERENCES_FILE = REPO_ROOT / "references.md"
ASSEMBLED_DIR = PROD_ROOT / "assembled"
METHOD_A_DIR = PROD_ROOT / "method_a_python"
METHOD_B_DIR = PROD_ROOT / "method_b_hybrid"
PROD_FIGURES_DIR = PROD_ROOT / "figures"

CONTENT_FILES = [
    "00_abstract.md",
    "01_introduction.md",
    "02_literature_review.md",
    "03_methodology.md",
    "04_results.md",
    "05_discussion.md",
    "06_recommendations_conclusion.md",
]


FRONT_MATTER_PAGES = [
    "Cover Page",
    "Certification of the Supervisor",
    "Dedication",
    "Acknowledgment",
    "Table of Contents",
    "List of Figures",
    "List of Tables",
    "List of Abbreviations",
]


CHAPTER_INSERTIONS = [
    ("# I. INTRODUCTION", "Chapter One", "Introduction"),
    ("# III. METHODOLOGY (ORIGINAL CROSS-SECTIONAL STUDY)", "Chapter Two", "Materials and Methods"),
    ("# IV. RESULTS", "Chapter Three", "Results"),
    ("# V. DISCUSSION", "Chapter Four", "Discussion"),
    ("# VI. RECOMMENDATIONS", "Chapter Five", "Conclusions and Suggestions"),
]


def ensure_dirs() -> None:
    for d in [ASSEMBLED_DIR, METHOD_A_DIR, METHOD_B_DIR, PROD_FIGURES_DIR]:
        d.mkdir(parents=True, exist_ok=True)


def copy_figure_assets() -> None:
    for png in FIGURES_DIR.glob("*.png"):
        shutil.copy2(png, PROD_FIGURES_DIR / png.name)
        shutil.copy2(png, ASSEMBLED_DIR / png.name)


def normalize_page_breaks(text: str) -> str:
    return text.replace(
        '<div style="page-break-after: always;"></div>',
        '<div class="page-break"></div>',
    )


def inject_chapter_title_pages(text: str) -> str:
    for heading, chapter_number, chapter_name in CHAPTER_INSERTIONS:
        marker = f"\n\n[[CHAPTER_TITLE:{chapter_number}|||{chapter_name}]]\n\n{heading}"
        text = text.replace(f"\n\n{heading}", marker, 1)
    return text


def inject_front_matter_pages(text: str) -> str:
    markers = []
    for page in FRONT_MATTER_PAGES:
        markers.append(f"[[FRONT_MATTER:{page}]]")
        markers.append('<div class="page-break"></div>')
    front_matter_block = "\n\n".join(markers).strip()
    return front_matter_block + "\n\n" + text


def transform_inline_figure_links(text: str) -> str:
    pattern = re.compile(r"\[([^\]]+\.(?:png|jpg|jpeg|webp))\]\(([^)]+\.(?:png|jpg|jpeg|webp))\)")
    out_lines: list[str] = []

    def friendly_label(raw_path: str) -> str:
        filename = Path(unquote(raw_path.strip())).name
        stem = Path(filename).stem
        return stem.replace("_", " ").strip() or "figure"

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line or line.lstrip().startswith("!"):
            out_lines.append(line)
            continue

        matches = list(pattern.finditer(line))
        if not matches:
            out_lines.append(line)
            continue

        cleaned = pattern.sub(lambda _: "the corresponding figure below", line)
        out_lines.append(cleaned)
        out_lines.append("")
        for m in matches:
            out_lines.append(f"![{friendly_label(m.group(2))}]({m.group(2)})")
        out_lines.append("")

    return "\n".join(out_lines)


def collect_figure_md_files() -> list[Path]:
    files = sorted(FIGURES_DIR.glob("*.md"), key=lambda p: p.name)
    return files


def rewrite_local_figure_links(md: str) -> str:
    def repl(match: re.Match[str]) -> str:
        alt = match.group(1)
        path = match.group(2)
        if "/" not in path:
            return f"![{alt}](../figures/{path})"
        return match.group(0)

    return re.sub(r"!\[([^\]]*)\]\(([^\)]+)\)", repl, md)


def assemble_markdown() -> Path:
    parts: list[str] = []

    for name in CONTENT_FILES:
        text = (CONTENT_DIR / name).read_text(encoding="utf-8")
        parts.append(text.strip())

    parts.append("# VIII. REFERENCES")
    parts.append(REFERENCES_FILE.read_text(encoding="utf-8").strip())

    for fig_md in collect_figure_md_files():
        fig_text = fig_md.read_text(encoding="utf-8")
        fig_text = rewrite_local_figure_links(fig_text)
        parts.append(fig_text.strip())

    merged = "\n\n".join(parts).strip() + "\n"
    merged = normalize_page_breaks(merged)
    merged = inject_front_matter_pages(merged)
    merged = inject_chapter_title_pages(merged)
    merged = transform_inline_figure_links(merged)

    out_path = ASSEMBLED_DIR / "comprehensive_research.md"
    out_path.write_text(merged, encoding="utf-8")
    shutil.copy2(out_path, METHOD_A_DIR / "comprehensive_research.md")
    shutil.copy2(out_path, METHOD_B_DIR / "comprehensive_research.md")
    return out_path


def iter_markdown_blocks(text: str) -> Iterable[tuple[str, str]]:
    paragraph_lines: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()

        if line.strip() == '<div class="page-break"></div>':
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("pagebreak", "")
            continue

        front_matter = re.match(r"^\[\[FRONT_MATTER:(.+)\]\]$", line.strip())
        if front_matter:
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("frontmatter", front_matter.group(1).strip())
            continue

        chapter_title = re.match(r"^\[\[CHAPTER_TITLE:(.+)\]\]$", line.strip())
        if chapter_title:
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("chaptertitle", chapter_title.group(1).strip())
            continue

        h1 = re.match(r"^#\s+(.+)$", line)
        h2 = re.match(r"^##\s+(.+)$", line)
        img = re.match(r"^!\[([^\]]*)\]\(([^\)]+)\)", line.strip())
        ordered_item = re.match(r"^\s*\d+\.\s+.+$", line)

        if h1:
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("h1", h1.group(1).strip())
            continue
        if h2:
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("h2", h2.group(1).strip())
            continue
        if img:
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("image", f"{img.group(1)}|||{img.group(2)}")
            continue
        if ordered_item:
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            yield ("paragraph", line.strip())
            continue

        if not line.strip():
            if paragraph_lines:
                yield ("paragraph", " ".join(paragraph_lines).strip())
                paragraph_lines = []
            continue

        paragraph_lines.append(line)

    if paragraph_lines:
        yield ("paragraph", " ".join(paragraph_lines).strip())


def build_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Inches(0.5)
    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Times New Roman"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_1.font.size = Pt(18)
    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Times New Roman"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_2.font.size = Pt(16)

    def add_centered_title_page(title: str, subtitle: str | None = None, with_rules: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Inches(3.2)
        p.paragraph_format.space_after = Inches(0.16)
        if with_rules:
            rule_top = p.add_run("────────────")
            rule_top.font.name = "Times New Roman"
            rule_top._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            rule_top.font.size = Pt(18)

        p2 = doc.add_paragraph(title)
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p2.paragraph_format.first_line_indent = Inches(0)
        p2.paragraph_format.space_after = Inches(0.08)
        run = p2.runs[0]
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(32 if with_rules else 28)
        run.bold = True

        if subtitle:
            p3 = doc.add_paragraph(subtitle)
            p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p3.paragraph_format.first_line_indent = Inches(0)
            p3.paragraph_format.space_after = Inches(0.08)
            sub_run = p3.runs[0]
            sub_run.font.name = "Times New Roman"
            sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            sub_run.font.size = Pt(18)
            sub_run.bold = True

        if with_rules:
            p4 = doc.add_paragraph()
            p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p4.paragraph_format.first_line_indent = Inches(0)
            rule_bottom = p4.add_run("────────────")
            rule_bottom.font.name = "Times New Roman"
            rule_bottom._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            rule_bottom.font.size = Pt(18)

        doc.add_page_break()

    started = False
    chapter_just_emitted = False
    in_references = False
    for kind, data in iter_markdown_blocks(text):
        if kind == "frontmatter":
            add_centered_title_page(data, with_rules=False)
            started = True
            chapter_just_emitted = False
            continue
        if kind == "chaptertitle":
            chapter_number, chapter_name = (data.split("|||", 1) + [""])[:2]
            add_centered_title_page(chapter_number, subtitle=chapter_name, with_rules=True)
            started = True
            chapter_just_emitted = True
            continue
        if kind == "h1":
            if started and not chapter_just_emitted:
                doc.add_page_break()
            p = doc.add_paragraph(data)
            p.style = doc.styles["Heading 1"]
            p.paragraph_format.first_line_indent = Inches(0)
            in_references = data.strip().upper() == "VIII. REFERENCES"
            started = True
            chapter_just_emitted = False
            continue
        if kind == "h2":
            p = doc.add_paragraph(data)
            p.style = doc.styles["Heading 2"]
            p.paragraph_format.first_line_indent = Inches(0)
            started = True
            chapter_just_emitted = False
            continue
        if kind == "paragraph":
            p = doc.add_paragraph(data)
            p.paragraph_format.line_spacing = 1.5
            if in_references:
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_after = Pt(8)
            else:
                p.paragraph_format.first_line_indent = Inches(0.5)
            started = True
            chapter_just_emitted = False
            continue
        if kind == "image":
            _, rel_path = data.split("|||", 1)
            img_path = (md_path.parent / rel_path).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.3))
            started = True
            chapter_just_emitted = False
            continue
        if kind == "pagebreak":
            doc.add_page_break()
            started = True
            chapter_just_emitted = False

    doc.save(out_path)


def escape_latex(text: str) -> str:
    table = {
        "\\": r"\\textbackslash{}",
        "&": r"\\&",
        "%": r"\\%",
        "$": r"\\$",
        "#": r"\\#",
        "_": r"\\_",
        "{": r"\\{",
        "}": r"\\}",
        "~": r"\\textasciitilde{}",
        "^": r"\\textasciicircum{}",
    }
    return "".join(table.get(ch, ch) for ch in text)


def build_tex(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    body: list[str] = []

    for kind, data in iter_markdown_blocks(text):
        if kind == "frontmatter":
            body.append("\\newpage")
            body.append("\\begin{center}\\vspace*{0.42\\textheight}")
            body.append("{\\Large \\textbf{" + escape_latex(data) + "}}")
            body.append("\\end{center}")
            body.append("\\newpage")
        elif kind == "chaptertitle":
            chapter_number, chapter_name = (data.split("|||", 1) + [""])[:2]
            body.append("\\newpage")
            body.append("\\begin{center}\\vspace*{0.42\\textheight}")
            body.append("{\\Large \\textbf{\\rule{5cm}{0.4pt}}}\\\\[0.4cm]")
            body.append("{\\LARGE \\textbf{" + escape_latex(chapter_number) + "}}\\\\[0.15cm]")
            body.append("{\\Large \\textbf{" + escape_latex(chapter_name) + "}}\\\\[0.35cm]")
            body.append("{\\Large \\textbf{\\rule{5cm}{0.4pt}}}")
            body.append("\\end{center}")
            body.append("\\newpage")
        elif kind == "h1":
            body.append("\\newpage")
            body.append(f"\\section*{{{escape_latex(data)}}}")
        elif kind == "h2":
            body.append(f"\\subsection*{{{escape_latex(data)}}}")
        elif kind == "paragraph":
            body.append(escape_latex(data) + "\n")
        elif kind == "image":
            alt, rel_path = data.split("|||", 1)
            rel = rel_path.replace('\\', '/')
            body.append(
                "\\begin{figure}[h!]\n"
                "\\centering\n"
                f"\\includegraphics[width=0.9\\textwidth]{{{escape_latex(rel)}}}\n"
                f"\\caption*{{{escape_latex(alt)}}}\n"
                "\\end{figure}"
            )
        elif kind == "pagebreak":
            body.append("\\newpage")

    tex = (
        "\\documentclass[12pt,a4paper]{article}\n"
        "\\usepackage[left=3cm,right=2cm,top=2cm,bottom=2cm]{geometry}\n"
        "\\usepackage{setspace}\n"
        "\\usepackage{graphicx}\n"
        "\\usepackage[T1]{fontenc}\n"
        "\\usepackage[utf8]{inputenc}\n"
        "\\usepackage{mathptmx}\n"
        "\\setstretch{1.5}\n"
        "\\setlength{\\parindent}{0.5in}\n"
        "\\begin{document}\n\n"
        + "\n\n".join(body)
        + "\n\n\\end{document}\n"
    )
    out_path.write_text(tex, encoding="utf-8")


def _fit_image_width(img_path: Path, max_width_cm: float = 16.0) -> float:
    try:
        img = ImageReader(str(img_path))
        iw, ih = img.getSize()
        if iw <= 0 or ih <= 0:
            return max_width_cm * cm
        ratio = ih / iw
        width_cm = max_width_cm
        height_cm = width_cm * ratio
        if height_cm > 20:
            scale = 20 / height_cm
            width_cm *= scale
        return width_cm * cm
    except Exception:
        return max_width_cm * cm


def build_pdf_reportlab(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=18,
        leading=24,
        spaceBefore=10,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=16,
        leading=22,
        spaceBefore=8,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=14,
        leading=21,
        firstLineIndent=0.5 * 72,
        spaceBefore=2,
        spaceAfter=4,
    )
    refs = ParagraphStyle(
        "Refs",
        parent=body,
        firstLineIndent=0,
        leftIndent=0,
        spaceBefore=1,
        spaceAfter=8,
    )

    story = []
    started = False
    chapter_just_emitted = False
    in_references = False
    chapter_num_style = ParagraphStyle(
        "ChapterNumber",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=34,
        leading=40,
        alignment=1,
    )
    chapter_name_style = ParagraphStyle(
        "ChapterName",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=18,
        leading=24,
        alignment=1,
    )
    front_matter_style = ParagraphStyle(
        "FrontMatterTitle",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=28,
        leading=34,
        alignment=1,
    )
    chapter_rule_style = ParagraphStyle(
        "ChapterRule",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=18,
        leading=22,
        alignment=1,
    )

    for kind, data in iter_markdown_blocks(text):
        if kind == "frontmatter":
            if started:
                story.append(PageBreak())
            story.append(Spacer(1, 260))
            story.append(Paragraph(data, front_matter_style))
            story.append(Spacer(1, 260))
            story.append(PageBreak())
            started = True
            chapter_just_emitted = False
        elif kind == "chaptertitle":
            chapter_number, chapter_name = (data.split("|||", 1) + [""])[:2]
            if started:
                story.append(PageBreak())
            story.append(Spacer(1, 220))
            story.append(Paragraph("────────────", chapter_rule_style))
            story.append(Spacer(1, 6))
            story.append(Paragraph(chapter_number, chapter_num_style))
            story.append(Spacer(1, 4))
            story.append(Paragraph(chapter_name, chapter_name_style))
            story.append(Spacer(1, 8))
            story.append(Paragraph("────────────", chapter_rule_style))
            story.append(Spacer(1, 220))
            story.append(PageBreak())
            started = True
            chapter_just_emitted = True
        elif kind == "h1":
            if started and not chapter_just_emitted:
                story.append(PageBreak())
            story.append(Paragraph(data, h1))
            in_references = data.strip().upper() == "VIII. REFERENCES"
            started = True
            chapter_just_emitted = False
        elif kind == "h2":
            story.append(Paragraph(data, h2))
            started = True
            chapter_just_emitted = False
        elif kind == "paragraph":
            story.append(Paragraph(data, refs if in_references else body))
            started = True
            chapter_just_emitted = False
        elif kind == "image":
            _, rel_path = data.split("|||", 1)
            img_path = (md_path.parent / rel_path).resolve()
            if img_path.exists():
                width = _fit_image_width(img_path)
                img_reader = ImageReader(str(img_path))
                iw, ih = img_reader.getSize()
                if iw > 0:
                    height = width * (ih / iw)
                else:
                    height = 8 * cm
                img = Image(str(img_path), width=width, height=height, hAlign="CENTER")
                story.append(Spacer(1, 8))
                story.append(img)
                story.append(Spacer(1, 8))
            started = True
            chapter_just_emitted = False
        elif kind == "pagebreak":
            story.append(PageBreak())
            started = True
            chapter_just_emitted = False

    doc.build(story)


def build_pdf_xhtml2pdf(md_path: Path, out_pdf: Path, out_html: Path) -> None:
    css = (PROD_ROOT / "templates" / "print.css").read_text(encoding="utf-8")
    md_text = md_path.read_text(encoding="utf-8")
    md_text = re.sub(
        r"\[\[FRONT_MATTER:(.+?)\]\]",
        r'<div class="page-break"></div><div class="front-matter-page">\1</div>',
        md_text,
    )
    md_text = re.sub(
        r"\[\[CHAPTER_TITLE:(.+?)\|\|\|(.+?)\]\]",
        r'<div class="page-break"></div><div class="chapter-title-page"><div class="chapter-rule">────────────</div><div class="chapter-title">\1</div><div class="chapter-subtitle">\2</div><div class="chapter-rule">────────────</div></div><div class="page-break"></div>',
        md_text,
    )
    html_body = markdown2.markdown(md_text, extras=["tables", "fenced-code-blocks", "cuddled-lists"])
    html_body = re.sub(
        r'(<img[^>]+src=")([^":]+?)(")',
        lambda m: (
            m.group(1)
            + str((md_path.parent / unquote(m.group(2))).resolve())
            + m.group(3)
        ),
        html_body,
    )
    html = (
        "<html><head><meta charset='utf-8'><style>"
        + css
        + "</style></head><body>"
        + html_body
        + "</body></html>"
    )

    out_html.write_text(html, encoding="utf-8")
    with out_pdf.open("wb") as f:
        pisa.CreatePDF(src=html, dest=f, path=str(md_path.parent))


def build_with_pandoc(md_path: Path, out_tex: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    md_text = re.sub(r"\[\[FRONT_MATTER:(.+?)\]\]", r"\n\n# \1\n\n<div class=\"page-break\"></div>\n\n", md_text)
    md_text = re.sub(
        r"\[\[CHAPTER_TITLE:(.+?)\|\|\|(.+?)\]\]",
        r"\n\n# \1\n\n## \2\n\n<div class=\"page-break\"></div>\n\n",
        md_text,
    )
    tmp_md = METHOD_B_DIR / "_pandoc_input.md"
    tmp_md.write_text(md_text, encoding="utf-8")
    pypandoc.convert_file(
        str(tmp_md),
        to="latex",
        format="md",
        outputfile=str(out_tex),
        extra_args=["--resource-path", str(md_path.parent)],
    )
    tmp_md.unlink(missing_ok=True)


def run_method_a(md_path: Path) -> None:
    build_docx(md_path, METHOD_A_DIR / "research_method_a.docx")
    build_tex(md_path, METHOD_A_DIR / "research_method_a.tex")
    build_pdf_reportlab(md_path, METHOD_A_DIR / "research_method_a.pdf")


def run_method_b(md_path: Path) -> None:
    build_docx(md_path, METHOD_B_DIR / "research_method_b.docx")
    build_with_pandoc(md_path, METHOD_B_DIR / "research_method_b.tex")
    build_pdf_xhtml2pdf(md_path, METHOD_B_DIR / "research_method_b.pdf", METHOD_B_DIR / "research_method_b.html")


def main() -> None:
    ensure_dirs()
    copy_figure_assets()
    md_path = assemble_markdown()
    run_method_a(md_path)
    run_method_b(md_path)
    print("Production pipeline completed.")
    print(f"Comprehensive markdown: {md_path}")
    print(f"Method A outputs: {METHOD_A_DIR}")
    print(f"Method B outputs: {METHOD_B_DIR}")


if __name__ == "__main__":
    main()
